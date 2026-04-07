"""Resume document conversion: Markdown → PDF and Markdown → DOCX.

PDF:  Uses Patchright Chromium (already a project dependency) with a custom
      HTML/CSS template tuned for single-column resumes. No date stamps, no
      file path headers, no page number footers. Clean ATS output.

DOCX: Uses python-docx with direct paragraph-level formatting. No pandoc
      dependency. Produces the same single-column layout as the PDF.

Expected resume Markdown format (produced by Huntsman):

    Name
    email | phone | location | github | linkedin
    email | phone [second contact line optional]

    PROFESSIONAL SUMMARY
    [paragraph]

    TECHNICAL SKILLS
    Category: item, item, item

    WORK EXPERIENCE
    Title | Company | Type | Dates
    - Bullet point
    - Bullet point
    - Tech: stack items

    PROJECTS
    Title | Year | URL
    - Bullet point
    - Tech: stack items

    EDUCATION
    Degree | School | Dates

    CERTIFICATIONS
    - Item

    AWARDS AND HACKATHONS
    - Item
"""

import asyncio
import re
from pathlib import Path, PurePosixPath

from huntsman_mcp.config import OUTPUT_DIR
from huntsman_mcp.exceptions import HuntsmanError


class ConversionError(HuntsmanError):
    """Document conversion failed."""


def _safe_filename(filename: str) -> str:
    name = PurePosixPath(filename).name
    if not name or name.startswith("."):
        raise ConversionError(
            f"Invalid filename {filename!r}. "
            "Provide a simple name like 'resume_fullstack' with no path separators."
        )
    return name


# LLMs frequently emit Markdown even when instructed not to. This sanitizer
# strips all formatting artifacts before any line-level processing so they
# never appear literally in the converted document.

_RE_FENCED_CODE = re.compile(r"```[^\n]*\n(.*?)```", re.DOTALL)
_RE_IMAGE = re.compile(r"!\[[^\]]*\]\([^)]*\)")
_RE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]*)\)")
_RE_BOLD_ITALIC = re.compile(r"\*{1,3}([^*\n]+)\*{1,3}")
_RE_UNDERLINE = re.compile(r"_{1,3}([^_\n]+)_{1,3}")
_RE_INLINE_CODE = re.compile(r"`+([^`\n]+)`+")
_RE_HTML_TAG = re.compile(r"<[^>]+>")
_RE_MD_HEADING = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_RE_HR = re.compile(r"^[-=*]{3,}\s*$", re.MULTILINE)


def _link_sub(m: re.Match) -> str:
    label, url = m.group(1), m.group(2)
    url_bare = re.sub(r"^https?://", "", url).rstrip("/")
    if url_bare == label or url_bare in label or label in url_bare:
        return label
    return f"{label} ({url})"


def _sanitize(text: str) -> str:
    text = _RE_FENCED_CODE.sub(lambda m: m.group(1).strip(), text)
    text = _RE_IMAGE.sub("", text)
    text = _RE_LINK.sub(_link_sub, text)
    text = _RE_BOLD_ITALIC.sub(r"\1", text)
    text = _RE_UNDERLINE.sub(r"\1", text)
    text = _RE_INLINE_CODE.sub(r"\1", text)
    text = _RE_HTML_TAG.sub("", text)
    text = _RE_MD_HEADING.sub("", text)
    text = _RE_HR.sub("", text)
    lines = [line.rstrip() for line in text.splitlines()]
    return "\n".join(lines)


_SECTION_HEADERS = frozenset(
    [
        "PROFESSIONAL SUMMARY",
        "TECHNICAL SKILLS",
        "WORK EXPERIENCE",
        "PROJECTS",
        "EDUCATION",
        "CERTIFICATIONS",
        "AWARDS AND HACKATHONS",
        "CERTIFICATIONS & AWARDS",
        "AWARDS",
        "HACKATHONS",
    ]
)


_CSS = """
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

body {
    font-family: 'Calibri', 'Arial', sans-serif;
    font-size: 10pt;
    line-height: 1.4;
    color: #111;
    background: #fff;
}

.name {
    font-size: 22pt;
    font-weight: 700;
    letter-spacing: 0.01em;
    color: #000;
    margin-bottom: 4px;
}

.contact {
    font-size: 8.5pt;
    color: #444;
    margin-bottom: 1px;
}

.section-head {
    font-size: 8.5pt;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #000;
    border-bottom: 1.2px solid #000;
    padding-bottom: 2px;
    margin-top: 13px;
    margin-bottom: 5px;
}

.entry-title {
    font-weight: 700;
    font-size: 9.5pt;
    color: #000;
    margin-top: 6px;
    margin-bottom: 2px;
}

.skill-row {
    font-size: 9pt;
    margin-bottom: 2px;
    line-height: 1.35;
}

.skill-cat { font-weight: 700; }

ul {
    margin-left: 15px;
    margin-bottom: 2px;
    margin-top: 1px;
}

li {
    font-size: 9pt;
    line-height: 1.4;
    margin-bottom: 1px;
    color: #111;
}

.tech {
    font-size: 8.5pt;
    color: #555;
    font-style: italic;
    margin-top: 1px;
    margin-bottom: 3px;
}

.body-line {
    font-size: 9pt;
    margin-bottom: 2px;
    line-height: 1.4;
}
"""


def _build_html(markdown_text: str) -> str:
    """Convert resume Markdown to HTML string for Chromium PDF rendering."""
    lines = markdown_text.strip().splitlines()
    parts: list[str] = []
    i = 0
    contact_count = 0

    # --- Name (first non-empty line) ---
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines):
        parts.append(f'<div class="name">{_escape(lines[i].strip())}</div>')
        i += 1

    # --- Contact lines (next 1-2 non-empty lines before a blank or section header) ---
    while i < len(lines) and contact_count < 2:
        stripped = lines[i].strip()
        if stripped and stripped.upper() not in _SECTION_HEADERS:
            parts.append(f'<div class="contact">{_escape(stripped)}</div>')
            contact_count += 1
            i += 1
        elif not stripped:
            i += 1
            break
        else:
            break

    # --- Body ---
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        upper = stripped.upper()

        # Section header
        if upper in _SECTION_HEADERS:
            parts.append(f'<div class="section-head">{_escape(stripped)}</div>')
            i += 1
            continue

        # Tech line (italic, smaller)
        if stripped.startswith("- Tech:"):
            parts.append(f'<div class="tech">{_escape(stripped[2:])}</div>')
            i += 1
            continue

        # Bullet list — collect consecutive bullets
        if stripped.startswith("- "):
            parts.append("<ul>")
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith("- ") and not s.startswith("- Tech:"):
                    parts.append(f"<li>{_escape(s[2:])}</li>")
                    i += 1
                else:
                    break
            parts.append("</ul>")
            continue

        # Skill category lines: "Category: items"
        if re.match(r"^[A-Za-z /]+:", stripped) and upper not in _SECTION_HEADERS:
            cat, _, items = stripped.partition(":")
            parts.append(
                f'<div class="skill-row">'
                f'<span class="skill-cat">{_escape(cat)}:</span>{_escape(items)}'
                f"</div>"
            )
            i += 1
            continue

        # Entry title (job, project, education — contains |)
        if "|" in stripped:
            parts.append(f'<div class="entry-title">{_escape(stripped)}</div>')
            i += 1
            continue

        # Regular body line
        parts.append(f'<div class="body-line">{_escape(stripped)}</div>')
        i += 1

    body = "\n".join(parts)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>{_CSS}</style>
</head>
<body>
{body}
</body>
</html>"""


def _escape(text: str) -> str:
    """Minimal HTML entity escaping — only what's needed for plain text content."""
    return (
        text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


async def to_pdf(markdown_text: str, filename: str) -> Path:
    if not filename.endswith(".pdf"):
        filename += ".pdf"
    filename = _safe_filename(filename)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / filename

    html = _build_html(_sanitize(markdown_text))

    from patchright.async_api import async_playwright

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            try:
                page = await browser.new_page()
                await page.set_content(html, wait_until="networkidle")
                await page.pdf(
                    path=str(output_path),
                    format="Letter",
                    print_background=True,
                    display_header_footer=False,
                    margin={
                        "top": "0.6in",
                        "bottom": "0.6in",
                        "left": "0.65in",
                        "right": "0.65in",
                    },
                )
            finally:
                await browser.close()
    except Exception as exc:
        raise ConversionError(f"PDF generation failed: {exc}") from exc

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise ConversionError(
            f"PDF was not created at {output_path}. Check that Patchright Chromium is installed."
        )

    return output_path


def to_docx(markdown_text: str, filename: str) -> Path:
    if not filename.endswith(".docx"):
        filename += ".docx"
    filename = _safe_filename(filename)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / filename

    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise ConversionError(
            "python-docx is required for DOCX conversion. "
            "Install it: pip install python-docx"
        ) from exc

    lines = _sanitize(markdown_text).strip().splitlines()
    doc = Document()

    # Remove default styles noise
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def _hr(d: "Document") -> None:
        """Add a horizontal rule beneath the current paragraph."""
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "111111")
        pBdr.append(bottom)
        pPr.append(pBdr)

    try:
        i = 0

        # --- Name ---
        while i < len(lines) and not lines[i].strip():
            i += 1
        if i < len(lines):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(lines[i].strip())
            run.bold = True
            run.font.size = Pt(16)
            i += 1

        # --- Contact lines (next 1-2 non-empty lines before a blank or section header) ---
        contact_count = 0
        while i < len(lines) and contact_count < 2:
            stripped = lines[i].strip()
            if stripped and stripped.upper() not in _SECTION_HEADERS:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                contact_count += 1
                i += 1
            elif not stripped:
                i += 1
                if contact_count > 0:
                    break
            else:
                break

        # --- Body ---
        while i < len(lines):
            stripped = lines[i].strip()
            upper = stripped.upper()

            if not stripped:
                i += 1
                continue

            # Section header
            if upper in _SECTION_HEADERS:
                _hr(doc)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
                i += 1
                continue

            # Tech line
            if stripped.startswith("- Tech:"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Inches(0.2)
                run = p.add_run(stripped[2:])
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                run.italic = True
                i += 1
                continue

            # Bullet
            if stripped.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.2)
                run = p.add_run(stripped[2:])
                run.font.size = Pt(10)
                i += 1
                continue

            # Skill category: "Category: items"
            if re.match(r"^[A-Za-z /]+:", stripped) and upper not in _SECTION_HEADERS:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                cat, _, items = stripped.partition(":")
                bold_run = p.add_run(cat + ":")
                bold_run.bold = True
                bold_run.font.size = Pt(9.5)
                normal_run = p.add_run(items)
                normal_run.font.size = Pt(9.5)
                i += 1
                continue

            # Entry title (contains |)
            if "|" in stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(10)
                i += 1
                continue

            # Regular line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            i += 1

        doc.save(str(output_path))

    except Exception as exc:
        raise ConversionError(f"DOCX generation failed: {exc}") from exc

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise ConversionError(f"DOCX was not created at {output_path}.")

    return output_path


async def to_docx_async(markdown_text: str, filename: str) -> Path:
    """Async wrapper around the synchronous DOCX converter.

    python-docx is synchronous. This offloads it to a thread pool so it
    doesn't block the MCP server's event loop.
    """
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, to_docx, markdown_text, filename)
